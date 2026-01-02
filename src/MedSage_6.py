"""
Enhanced MedSage Medical RAG System
- Fixed all output errors and user interaction issues
- Added comprehensive error handling
- Enhanced gender-specific modules with user interaction
- Added medication reminders, appointment scheduling, and health insights
- Improved MongoDB operations with proper error handling
- Complete Ollama integration with fallback responses
"""

import os #For interaction with files & directories and environment variables
import re #For regular expressions to validate email formats through pattern matching, validation, and search and replace operations
import json #To allow serilisation and deserialization of data in JSON format
import shutil #Handling file operations like copying and moving files
import hashlib #For password hashing 
import uuid #To generate unique identifiers for users, sessions, and records
import getpass #To securely handle password input from users without echoing
import warnings #To manage and suppress warnings during execution
from pathlib import Path #For handling and manipulating filesystem paths
from typing import List, Dict, Optional, Any, Tuple #For type hinting to improve code clarity and maintainability
from datetime import datetime, timedelta #To handle date and time operations
from collections import defaultdict #To create dictionaries with default values for easier data aggregation

import pandas as pd #For data manipulation and analysis
import numpy as np #For numerical operations and array handling

warnings.filterwarnings("ignore")

# Optional dependencies with graceful fallbacks
try:
    from pymongo import MongoClient, ASCENDING, DESCENDING
    from pymongo.errors import DuplicateKeyError, ConnectionFailure
    HAS_MONGODB = True
except ImportError:
    MongoClient = None
    ASCENDING = DESCENDING = None
    DuplicateKeyError = ConnectionFailure = Exception
    HAS_MONGODB = False

try:
    import PyPDF2
    HAS_PDF = True
except ImportError:
    HAS_PDF = False

try:
    from PIL import Image
    import pytesseract
    HAS_OCR = True
except ImportError:
    HAS_OCR = False

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

try:
    import matplotlib
    matplotlib.use("Agg")
    import matplotlib.pyplot as plt
    import seaborn as sns
    HAS_PLOTTING = True
except ImportError:
    HAS_PLOTTING = False

try:
    from langchain_community.llms import Ollama
    from langchain_community.embeddings import HuggingFaceEmbeddings
    from langchain_community.vectorstores import Chroma
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    from langchain.schema import Document as LangChainDocument
    HAS_LANGCHAIN = True
except ImportError:
    HAS_LANGCHAIN = False
    class LangChainDocument:
        def __init__(self, page_content: str, metadata: Dict = None):
            self.page_content = page_content
            self.metadata = metadata or {}

# Utilities
def validate_email(email: str) -> bool:
    """Validate email format"""
    pattern = r'^[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}$'
    return bool(re.match(pattern, (email or "").strip()))

def normalize_email(email: str) -> str:
    """Normalize email to lowercase"""
    return (email or "").strip().lower()

def safe_input(prompt: str, default: str = "") -> str:
    """Safe input with default fallback"""
    try:
        value = input(prompt).strip()
        return value if value else default
    except (EOFError, KeyboardInterrupt):
        return default

def safe_int_input(prompt: str, default: int = 0, min_val: int = None, max_val: int = None) -> int:
    """Safe integer input with validation"""
    try:
        value = input(prompt).strip()
        if not value:
            return default
        num = int(value)
        if min_val is not None and num < min_val:
            return min_val
        if max_val is not None and num > max_val:
            return max_val
        return num
    except (ValueError, EOFError, KeyboardInterrupt):
        return default

def safe_float_input(prompt: str, default: float = 0.0, min_val: float = None, max_val: float = None) -> float:
    """Safe float input with validation"""
    try:
        value = input(prompt).strip()
        if not value:
            return default
        num = float(value)
        if min_val is not None and num < min_val:
            return min_val
        if max_val is not None and num > max_val:
            return max_val
        return num
    except (ValueError, EOFError, KeyboardInterrupt):
        return default


class LifestyleTracker:
    """Enhanced lifestyle tracking with better error handling"""
    
    def __init__(self, db):
        self.db = db
        self.lifestyle = None
        self.water = None
        self.meals = None
        
        if db is not None:
            try:
                self.lifestyle = db.lifestyle_data
                self.water = db.water_intake
                self.meals = db.meal_logs
                
                # Create indexes
                self.lifestyle.create_index([("user_id", ASCENDING), ("date", DESCENDING)])
                self.water.create_index([("user_id", ASCENDING), ("date", DESCENDING)])
                self.meals.create_index([("user_id", ASCENDING), ("date", DESCENDING)])
            except Exception as e:
                print(f"Warning: Could not initialize lifestyle tracker: {e}")

    def log_daily_lifestyle(self, user_id: str, date: str, data: Dict) -> Dict:
        """Log daily lifestyle data with validation"""
        try:
            if self.lifestyle is None:
                return {"success": False, "error": "Database not available"}
            
            # Validate data
            required_fields = ["sleep_hours", "exercise_minutes", "mood", "stress_level"]
            for field in required_fields:
                if field not in data:
                    return {"success": False, "error": f"Missing field: {field}"}
            
            entry_id = str(uuid.uuid4())
            doc = {
                "entry_id": entry_id,
                "user_id": user_id,
                "date": date,
                "data": data,
                "logged_at": datetime.now()
            }
            
            # Update if exists, insert if new
            self.lifestyle.update_one(
                {"user_id": user_id, "date": date},
                {"$set": doc},
                upsert=True
            )
            
            return {"success": True, "entry_id": entry_id, "message": "Lifestyle data logged successfully"}
        except Exception as e:
            return {"success": False, "error": f"Failed to log lifestyle data: {e}"}

    def get_lifestyle_summary(self, user_id: str, days: int = 7) -> Dict:
        """Get lifestyle summary with proper error handling"""
        try:
            if not self.lifestyle:
                return {"success": False, "error": "Database not available"}
            
            cutoff = (datetime.now() - timedelta(days=days)).date().isoformat()
            entries = list(self.lifestyle.find(
                {"user_id": user_id, "date": {"$gte": cutoff}},
                {"_id": 0}
            ).sort("date", ASCENDING))
            
            if not entries:
                return {
                    "success": True,
                    "message": "No lifestyle data available for this period",
                    "period_days": days,
                    "total_entries": 0
                }
            
            # Extract metrics
            sleep = [e["data"].get("sleep_hours") for e in entries if e["data"].get("sleep_hours") is not None]
            exercise = [e["data"].get("exercise_minutes") for e in entries if e["data"].get("exercise_minutes") is not None]
            mood = [e["data"].get("mood") for e in entries if e["data"].get("mood") is not None]
            stress = [e["data"].get("stress_level") for e in entries if e["data"].get("stress_level") is not None]
            
            return {
                "success": True,
                "period_days": days,
                "total_entries": len(entries),
                "avg_sleep_hours": round(float(np.mean(sleep)), 1) if sleep else None,
                "avg_exercise_minutes": round(float(np.mean(exercise)), 1) if exercise else None,
                "avg_mood": round(float(np.mean(mood)), 1) if mood else None,
                "avg_stress": round(float(np.mean(stress)), 1) if stress else None,
                "entries": entries
            }
        except Exception as e:
            return {"success": False, "error": f"Failed to get lifestyle summary: {e}"}

    def log_water_intake(self, user_id: str, date: str, amount_ml: int, time: str = None) -> Dict:
        """Log water intake with validation"""
        try:
            if self.water is None:
                return {"success": False, "error": "Database not available"}
            
            if amount_ml <= 0:
                return {"success": False, "error": "Amount must be positive"}
            
            entry_id = str(uuid.uuid4())
            doc = {
                "entry_id": entry_id,
                "user_id": user_id,
                "date": date,
                "time": time or datetime.now().strftime("%H:%M"),
                "amount_ml": amount_ml,
                "logged_at": datetime.now()
            }
            
            self.water.insert_one(doc)
            return {"success": True, "entry_id": entry_id, "message": "Water intake logged"}
        except Exception as e:
            return {"success": False, "error": f"Failed to log water: {e}"}

    def get_daily_water_intake(self, user_id: str, date: str) -> Dict:
        """Get daily water intake summary"""
        try:
            if not self.water:
                return {"success": False, "error": "Database not available"}
            
            entries = list(self.water.find({"user_id": user_id, "date": date}, {"_id": 0}))
            total = sum(e.get("amount_ml", 0) for e in entries)
            target = 2000
            
            return {
                "success": True,
                "date": date,
                "total_ml": total,
                "target_ml": target,
                "percentage": round((total / target) * 100, 1) if target else 0,
                "entries": entries,
                "glasses_8oz": round(total / 237)
            }
        except Exception as e:
            return {"success": False, "error": f"Failed to get water intake: {e}"}

    def log_meal(self, user_id: str, date: str, meal_type: str, description: str, calories: int = None) -> Dict:
        """Log meal information"""
        try:
            if self.meals is None:
                return {"success": False, "error": "Database not available"}
            
            meal_id = str(uuid.uuid4())
            doc = {
                "meal_id": meal_id,
                "user_id": user_id,
                "date": date,
                "meal_type": meal_type.lower(),
                "description": description,
                "calories": calories,
                "logged_at": datetime.now()
            }
            
            self.meals.insert_one(doc)
            return {"success": True, "meal_id": meal_id, "message": "Meal logged successfully"}
        except Exception as e:
            return {"success": False, "error": f"Failed to log meal: {e}"}


class WomensHealthModule:
    """Enhanced women's health tracking"""
    
    def __init__(self, db):
        self.db = db
        self.periods = None
        self.symptoms = None
        
        if db is not None:
            try:
                self.periods = db.womens_periods
                self.symptoms = db.womens_symptoms
                self.periods.create_index([("user_id", ASCENDING), ("start_date", DESCENDING)])
                self.symptoms.create_index([("user_id", ASCENDING), ("date", DESCENDING)])
            except Exception as e:
                print(f"Warning: Could not initialize women's health module: {e}")

    def log_period(self, user_id: str, start_date: str, end_date: Optional[str] = None, 
                   flow: str = "medium", notes: str = "") -> Dict:
        """Log menstrual period"""
        try:
            if self.periods is None:
                return {"success": False, "error": "Database not available"}
            
            period_id = str(uuid.uuid4())
            doc = {
                "period_id": period_id,
                "user_id": user_id,
                "start_date": start_date,
                "end_date": end_date,
                "flow": flow,
                "notes": notes,
                "logged_at": datetime.now()
            }
            
            if end_date:
                try:
                    s = datetime.fromisoformat(start_date)
                    e = datetime.fromisoformat(end_date)
                    doc["duration_days"] = (e - s).days + 1
                except Exception:
                    doc["duration_days"] = None
            
            self.periods.insert_one(doc)
            return {"success": True, "period_id": period_id, "message": "Period logged successfully"}
        except Exception as e:
            return {"success": False, "error": f"Failed to log period: {e}"}

    def predict_next_period(self, user_id: str) -> Dict:
        """Predict next period based on history"""
        try:
            if not self.periods:
                return {"success": False, "error": "Database not available"}
            
            periods = list(self.periods.find(
                {"user_id": user_id},
                {"_id": 0}
            ).sort("start_date", DESCENDING).limit(6))
            
            if len(periods) < 2:
                return {"success": False, "error": "Need at least 2 period entries for prediction"}
            
            cycle_lengths = []
            for i in range(len(periods) - 1):
                try:
                    d1 = datetime.fromisoformat(periods[i]["start_date"])
                    d2 = datetime.fromisoformat(periods[i + 1]["start_date"])
                    cycle_lengths.append((d1 - d2).days)
                except Exception:
                    pass
            
            if not cycle_lengths:
                return {"success": False, "error": "Insufficient valid dates"}
            
            avg = float(np.mean(cycle_lengths))
            std = float(np.std(cycle_lengths))
            last = datetime.fromisoformat(periods[0]["start_date"])
            next_period = last + timedelta(days=int(avg))
            ovulation = next_period - timedelta(days=14)
            fertile_start = ovulation - timedelta(days=5)
            fertile_end = ovulation + timedelta(days=1)
            
            return {
                "success": True,
                "next_period_date": next_period.date().isoformat(),
                "confidence": "high" if std < 3 else "medium" if std < 5 else "low",
                "average_cycle_length": round(avg, 1),
                "cycle_variation": round(std, 1),
                "ovulation_date": ovulation.date().isoformat(),
                "fertile_window_start": fertile_start.date().isoformat(),
                "fertile_window_end": fertile_end.date().isoformat()
            }
        except Exception as e:
            return {"success": False, "error": f"Failed to predict period: {e}"}

    def log_symptom(self, user_id: str, date: str, symptom_type: str, severity: int, notes: str = "") -> Dict:
        """Log health symptoms"""
        try:
            if self.symptoms is None:
                return {"success": False, "error": "Database not available"}
            
            symptom_id = str(uuid.uuid4())
            doc = {
                "symptom_id": symptom_id,
                "user_id": user_id,
                "date": date,
                "symptom_type": symptom_type,
                "severity": severity,
                "notes": notes,
                "logged_at": datetime.now()
            }
            
            self.symptoms.insert_one(doc)
            return {"success": True, "symptom_id": symptom_id, "message": "Symptom logged"}
        except Exception as e:
            return {"success": False, "error": f"Failed to log symptom: {e}"}


class MensHealthModule:
    """Enhanced men's health tracking"""
    
    def __init__(self, db):
        self.db = db
        self.health = None
        self.fitness = None
        
        if db is not None:
            try:
                self.health = db.mens_health
                self.fitness = db.mens_fitness
                self.health.create_index([("user_id", ASCENDING), ("date", DESCENDING)])
                self.fitness.create_index([("user_id", ASCENDING), ("date", DESCENDING)])
            except Exception as e:
                print(f"Warning: Could not initialize men's health module: {e}")

    def log_daily_health(self, user_id: str, date: str, data: Dict) -> Dict:
        """Log daily health metrics"""
        try:
            if self.health is None:
                return {"success": False, "error": "Database not available"}
            
            log_id = str(uuid.uuid4())
            doc = {
                "log_id": log_id,
                "user_id": user_id,
                "date": date,
                **data,
                "logged_at": datetime.now()
            }
            
            self.health.update_one(
                {"user_id": user_id, "date": date},
                {"$set": doc},
                upsert=True
            )
            
            return {"success": True, "log_id": log_id, "message": "Health data logged"}
        except Exception as e:
            return {"success": False, "error": f"Failed to log health data: {e}"}

    def get_health_insights(self, user_id: str, days: int = 30) -> Dict:
        """Get health insights"""
        try:
            if not self.health:
                return {"success": False, "error": "Database not available"}
            
            cutoff = (datetime.now() - timedelta(days=days)).date().isoformat()
            entries = list(self.health.find(
                {"user_id": user_id, "date": {"$gte": cutoff}},
                {"_id": 0}
            ).sort("date", 1))
            
            if not entries:
                return {
                    "success": True,
                    "message": "No health data available",
                    "period_days": days,
                    "total_entries": 0
                }
            
            energy = [e.get("energy_level") for e in entries if e.get("energy_level") is not None]
            sleep_q = [e.get("sleep_quality") for e in entries if e.get("sleep_quality") is not None]
            
            return {
                "success": True,
                "period_days": days,
                "total_entries": len(entries),
                "avg_energy": round(float(np.mean(energy)), 1) if energy else None,
                "avg_sleep_quality": round(float(np.mean(sleep_q)), 1) if sleep_q else None,
                "trend": "stable"
            }
        except Exception as e:
            return {"success": False, "error": f"Failed to get insights: {e}"}


class MedicationManager:
    """Medication tracking and reminders"""
    
    def __init__(self, db):
        self.db = db
        self.medications = None
        self.doses = None
        
        if db is not None:
            try:
                self.medications = db.medications
                self.doses = db.medication_doses
                self.medications.create_index([("user_id", ASCENDING), ("is_active", DESCENDING)])
                self.doses.create_index([("user_id", ASCENDING), ("date", DESCENDING)])
            except Exception as e:
                print(f"Warning: Could not initialize medication manager: {e}")

    def add_medication(self, user_id: str, name: str, dosage: str, frequency: str, 
                      start_date: str, end_date: str = None, notes: str = "") -> Dict:
        """Add a new medication"""
        try:
            if self.medications is None:
                return {"success": False, "error": "Database not available"}
            
            med_id = str(uuid.uuid4())
            doc = {
                "medication_id": med_id,
                "user_id": user_id,
                "name": name,
                "dosage": dosage,
                "frequency": frequency,
                "start_date": start_date,
                "end_date": end_date,
                "notes": notes,
                "is_active": True,
                "added_at": datetime.now()
            }
            
            self.medications.insert_one(doc)
            return {"success": True, "medication_id": med_id, "message": "Medication added"}
        except Exception as e:
            return {"success": False, "error": f"Failed to add medication: {e}"}

    def log_dose(self, user_id: str, medication_id: str, date: str, time: str, taken: bool = True) -> Dict:
        """Log medication dose"""
        try:
            if self.doses is None:
                return {"success": False, "error": "Database not available"}
            
            dose_id = str(uuid.uuid4())
            doc = {
                "dose_id": dose_id,
                "user_id": user_id,
                "medication_id": medication_id,
                "date": date,
                "time": time,
                "taken": taken,
                "logged_at": datetime.now()
            }
            
            self.doses.insert_one(doc)
            return {"success": True, "dose_id": dose_id, "message": "Dose logged"}
        except Exception as e:
            return {"success": False, "error": f"Failed to log dose: {e}"}

    def get_active_medications(self, user_id: str) -> Dict:
        """Get all active medications"""
        try:
            if not self.medications:
                return {"success": False, "error": "Database not available"}
            
            meds = list(self.medications.find(
                {"user_id": user_id, "is_active": True},
                {"_id": 0}
            ))
            
            return {"success": True, "count": len(meds), "medications": meds}
        except Exception as e:
            return {"success": False, "error": f"Failed to get medications: {e}"}


class AppointmentScheduler:
    """Appointment management system"""
    
    def __init__(self, db):
        self.db = db
        self.appointments = None
        
        if db is not None:
            try:
                self.appointments = db.appointments
                self.appointments.create_index([("user_id", ASCENDING), ("date", ASCENDING)])
            except Exception as e:
                print(f"Warning: Could not initialize appointment scheduler: {e}")

    def schedule_appointment(self, user_id: str, date: str, time: str, doctor: str, 
                           purpose: str, location: str = "", notes: str = "") -> Dict:
        """Schedule a new appointment"""
        try:
            if self.appointments is None:
                return {"success": False, "error": "Database not available"}
            
            appt_id = str(uuid.uuid4())
            doc = {
                "appointment_id": appt_id,
                "user_id": user_id,
                "date": date,
                "time": time,
                "doctor": doctor,
                "purpose": purpose,
                "location": location,
                "notes": notes,
                "status": "scheduled",
                "created_at": datetime.now()
            }
            
            self.appointments.insert_one(doc)
            return {"success": True, "appointment_id": appt_id, "message": "Appointment scheduled"}
        except Exception as e:
            return {"success": False, "error": f"Failed to schedule appointment: {e}"}

    def get_upcoming_appointments(self, user_id: str, days: int = 30) -> Dict:
        """Get upcoming appointments"""
        try:
            if not self.appointments:
                return {"success": False, "error": "Database not available"}
            
            today = datetime.now().date().isoformat()
            future = (datetime.now() + timedelta(days=days)).date().isoformat()
            
            appts = list(self.appointments.find(
                {
                    "user_id": user_id,
                    "date": {"$gte": today, "$lte": future},
                    "status": "scheduled"
                },
                {"_id": 0}
            ).sort("date", ASCENDING))
            
            return {"success": True, "count": len(appts), "appointments": appts}
        except Exception as e:
            return {"success": False, "error": f"Failed to get appointments: {e}"}


class MedicalFileProcessor:
    """Enhanced file processing with better error handling"""
    
    def __init__(self, upload_dir: str = "./uploads"):
        self.upload_dir = Path(upload_dir)
        self.upload_dir.mkdir(parents=True, exist_ok=True)
        self.supported_extensions = {
            "pdf": [".pdf"],
            "images": [".jpg", ".jpeg", ".png", ".bmp", ".tiff"],
            "text": [".txt"],
            "docx": [".docx", ".doc"]
        }

    def _is_supported(self, ext: str) -> bool:
        return any(ext.lower() in vals for vals in self.supported_extensions.values())

    def save_uploaded_file(self, user_id: str, file_path: str, report_type: str) -> Dict:
        """Save and process uploaded file"""
        try:
            src = Path(file_path)
            if not src.exists():
                return {"success": False, "error": "File not found"}
            
            ext = src.suffix.lower()
            if not self._is_supported(ext):
                return {"success": False, "error": f"Unsupported file type: {ext}"}
            
            user_dir = self.upload_dir / user_id
            user_dir.mkdir(parents=True, exist_ok=True)
            
            file_id = str(uuid.uuid4())[:8]
            dest = user_dir / f"{file_id}_{src.name}"
            shutil.copy2(src, dest)
            
            extracted_text = self.extract_text(dest)
            
            return {
                "success": True,
                "file_id": file_id,
                "file_path": str(dest),
                "original_name": src.name,
                "file_type": ext,
                "extracted_text": extracted_text,
                "file_size": dest.stat().st_size
            }
        except Exception as e:
            return {"success": False, "error": f"File processing failed: {e}"}

    def extract_text(self, file_path: Path) -> str:
        """Extract text from various file formats"""
        ext = file_path.suffix.lower()
        try:
            if ext == ".pdf":
                return self._extract_from_pdf(file_path)
            elif ext in self.supported_extensions["images"]:
                return self._extract_from_image(file_path)
            elif ext == ".txt":
                return self._extract_from_text(file_path)
            elif ext in self.supported_extensions["docx"]:
                return self._extract_from_docx(file_path)
            return ""
        except Exception as e:
            return f"[Error extracting text: {e}]"

    def _extract_from_pdf(self, path: Path) -> str:
        if not HAS_PDF:
            return "[PDF support not installed - install PyPDF2]"
        try:
            text_pages = []
            with open(path, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text() or ""
                    text_pages.append(text)
            return "\n\n".join(text_pages)
        except Exception as e:
            return f"[PDF extraction error: {e}]"

    def _extract_from_image(self, path: Path) -> str:
        if not HAS_OCR:
            return "[OCR support not installed - install pillow and pytesseract]"
        try:
            img = Image.open(path)
            return pytesseract.image_to_string(img)
        except Exception as e:
            return f"[OCR error: {e}]"

    def _extract_from_text(self, path: Path) -> str:
        try:
            return path.read_text(encoding="utf-8")
        except Exception as e:
            return f"[Text read error: {e}]"

    def _extract_from_docx(self, path: Path) -> str:
        if not HAS_DOCX:
            return "[DOCX support not installed - install python-docx]"
        try:
            doc = docx.Document(path)
            return "\n".join(p.text for p in doc.paragraphs)
        except Exception as e:
            return f"[DOCX error: {e}]"


class MedicalReportManager:
    """Enhanced report management"""
    
    def __init__(self, db, file_processor: MedicalFileProcessor):
        self.db = db
        self.fp = file_processor
        self.reports = None
        
        if db is not None:
            try:
                self.reports = db.medical_reports
                self.reports.create_index([("user_id", ASCENDING), ("upload_date", DESCENDING)])
            except Exception as e:
                print(f"Warning: Could not initialize report manager: {e}")

    def upload_report(self, user_id: str, file_path: str, report_type: str, 
                     report_name: str, notes: str = "") -> Dict:
        """Upload and process medical report"""
        try:
            if self.reports is None:
                return {"success": False, "error": "Database not available"}
            
            res = self.fp.save_uploaded_file(user_id, file_path, report_type)
            if not res.get("success"):
                return res
            
            report_id = str(uuid.uuid4())
            doc = {
                "report_id": report_id,
                "user_id": user_id,
                "file_id": res["file_id"],
                "report_name": report_name,
                "report_type": report_type,
                "original_filename": res["original_name"],
                "file_path": res["file_path"],
                "file_type": res["file_type"],
                "file_size": res["file_size"],
                "extracted_text": res["extracted_text"],
                "upload_date": datetime.now(),
                "notes": notes,
                "tags": []
            }
            
            self.reports.insert_one(doc)
            
            return {
                "success": True,
                "report_id": report_id,
                "message": "Report uploaded successfully",
                "extracted_text_length": len(res["extracted_text"] or "")
            }
        except Exception as e:
            return {"success": False, "error": f"Report upload failed: {e}"}

    def get_user_reports(self, user_id: str, limit: int = 50) -> List[Dict]:
        """Get user's medical reports"""
        try:
            if not self.reports:
                return []
            return list(self.reports.find(
                {"user_id": user_id},
                {"_id": 0}
            ).sort("upload_date", DESCENDING).limit(limit))
        except Exception as e:
            print(f"Error fetching reports: {e}")
            return []

    def get_report_content(self, report_id: str) -> Optional[Dict]:
        """Get specific report content"""
        try:
            if not self.reports:
                return None
            return self.reports.find_one({"report_id": report_id}, {"_id": 0})
        except Exception as e:
            print(f"Error fetching report: {e}")
            return None

    def delete_report(self, report_id: str) -> Dict:
        """Delete a medical report"""
        try:
            if not self.reports:
                return {"success": False, "error": "Database not available"}
            
            report = self.reports.find_one({"report_id": report_id})
            if not report:
                return {"success": False, "error": "Report not found"}
            
            try:
                Path(report.get("file_path", "")).unlink(missing_ok=True)
            except Exception:
                pass
            
            self.reports.delete_one({"report_id": report_id})
            return {"success": True, "message": "Report deleted"}
        except Exception as e:
            return {"success": False, "error": f"Failed to delete report: {e}"}


class CSVDataProcessor:
    """Process CSV medical data files"""
    
    def __init__(self, data_dir: str):
        self.data_dir = Path(data_dir)
        self.data_dir.mkdir(parents=True, exist_ok=True)

    def process_all_files(self) -> List[LangChainDocument]:
        """Process all CSV files in data directory"""
        documents = []
        csv_files = list(self.data_dir.glob("**/*.csv"))
        
        print(f"Processing {len(csv_files)} CSV files...")
        
        for csv_file in csv_files:
            try:
                df = pd.read_csv(csv_file, low_memory=False)
                df.columns = df.columns.str.strip().str.lower()
                
                for idx, row in df.iterrows():
                    content = " | ".join([
                        f"{col}: {row[col]}" 
                        for col in df.columns 
                        if pd.notna(row[col])
                    ])
                    
                    if content:
                        documents.append(LangChainDocument(
                            page_content=content,
                            metadata={"source": str(csv_file), "row": int(idx)}
                        ))
            except Exception as e:
                print(f"Error processing {csv_file}: {e}")
        
        print(f"Processed {len(documents)} documents")
        return documents


class MedicalVectorStore:
    """Vector store for medical knowledge"""
    
    def __init__(self, persist_directory: str):
        self.persist_directory = persist_directory
        
        if HAS_LANGCHAIN:
            try:
                self.embeddings = HuggingFaceEmbeddings(
                    model_name="sentence-transformers/all-MiniLM-L6-v2",
                    model_kwargs={"device": "cpu"}
                )
                self.text_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1000,
                    chunk_overlap=200
                )
                print("Vector store initialized with embeddings")
            except Exception as e:
                print(f"Error initializing embeddings: {e}")
                self.embeddings = None
                self.text_splitter = None
        else:
            self.embeddings = None
            self.text_splitter = None

    def create_vectorstore(self, documents: List[LangChainDocument]):
        """Create new vector store from documents"""
        if not HAS_LANGCHAIN or not self.embeddings:
            raise RuntimeError("LangChain components not available")
        
        print("Creating vector store...")
        chunks = self.text_splitter.split_documents(documents)
        print(f"Split into {len(chunks)} chunks")
        
        vectorstore = Chroma.from_documents(
            documents=chunks,
            embedding=self.embeddings,
            persist_directory=self.persist_directory
        )
        print("Vector store created successfully")
        return vectorstore

    def load_vectorstore(self):
        """Load existing vector store"""
        if not HAS_LANGCHAIN or not self.embeddings:
            return None
        
        try:
            return Chroma(
                persist_directory=self.persist_directory,
                embedding_function=self.embeddings
            )
        except Exception as e:
            print(f"Error loading vector store: {e}")
            return None


class UserManager:
    """Comprehensive user management system"""
    
    def __init__(self, connection_string: str = "mongodb://localhost:27017/", 
                 database_name: str = "medical_rag_system"):
        if not HAS_MONGODB:
            raise ImportError("pymongo is required. Install with: pip install pymongo")
        
        try:
            self.client = MongoClient(connection_string, serverSelectionTimeoutMS=5000)
            self.client.server_info()
            print("✓ Connected to MongoDB")
        except Exception as e:
            raise ConnectionError(f"Failed to connect to MongoDB: {e}")
        
        self.db = self.client[database_name]
        self.users = self.db.users
        self.profiles = self.db.user_profiles
        self.conversations = self.db.conversations
        
        # Remove legacy username index if exists
        try:
            indexes = self.users.index_information()
            if "username_1" in indexes:
                self.users.drop_index("username_1")
        except Exception:
            pass
        
        # Initialize all modules
        self.file_processor = MedicalFileProcessor()
        self.report_manager = MedicalReportManager(self.db, self.file_processor)
        self.lifestyle_tracker = LifestyleTracker(self.db)
        self.womens_health = WomensHealthModule(self.db)
        self.mens_health = MensHealthModule(self.db)
        self.medication_manager = MedicationManager(self.db)
        self.appointment_scheduler = AppointmentScheduler(self.db)
        
        # Create indexes
        try:
            self.users.create_index([("email", ASCENDING)], unique=True)
            self.profiles.create_index([("user_id", ASCENDING)], unique=True)
            self.conversations.create_index([("user_id", ASCENDING), ("timestamp", DESCENDING)])
            print("✓ Database indexes created")
        except Exception as e:
            print(f"Warning: Index creation issue: {e}")

    def _hash_password(self, password: str) -> str:
        """Hash password using SHA256"""
        return hashlib.sha256(password.encode()).hexdigest()

    def signup(self, email: str, password: str, full_name: str, age: int, gender: str) -> Dict:
        """Register new user"""
        try:
            email = normalize_email(email)
            
            if not validate_email(email):
                return {"success": False, "error": "Invalid email format"}
            
            # Check if email already exists
            if self.users.find_one({"email": email}):
                return {"success": False, "error": "Email already registered"}
            
            if len(password) < 6:
                return {"success": False, "error": "Password must be at least 6 characters"}
            
            user_id = str(uuid.uuid4())
            
            user_doc = {
                "user_id": user_id,
                "email": email,
                "password_hash": self._hash_password(password),
                "full_name": full_name,
                "created_at": datetime.now(),
                "is_active": True
            }
            
            self.users.insert_one(user_doc)
            self._create_profile(user_id, age, gender)
            
            return {
                "success": True,
                "user_id": user_id,
                "message": "Registration successful! Please login."
            }
        except DuplicateKeyError:
            return {"success": False, "error": "Email already registered"}
        except Exception as e:
            return {"success": False, "error": f"Registration failed: {e}"}

    def login(self, email: str, password: str) -> Dict:
        """Authenticate user"""
        try:
            email = normalize_email(email)
            
            if not validate_email(email):
                return {"success": False, "error": "Invalid email format"}
            
            user = self.users.find_one({"email": email})
            if not user:
                return {"success": False, "error": "Email not found"}
            
            if user.get("password_hash") != self._hash_password(password):
                return {"success": False, "error": "Invalid password"}
            
            if not user.get("is_active", True):
                return {"success": False, "error": "Account is inactive"}
            
            # Update last login
            self.users.update_one(
                {"user_id": user["user_id"]},
                {"$set": {"last_login": datetime.now()}}
            )
            
            return {
                "success": True,
                "user_id": user["user_id"],
                "email": user["email"],
                "full_name": user.get("full_name", ""),
                "message": "Login successful!"
            }
        except Exception as e:
            return {"success": False, "error": f"Login failed: {e}"}

    def _create_profile(self, user_id: str, age: int, gender: str):
        """Create user profile"""
        profile = {
            "user_id": user_id,
            "personal_info": {
                "age": age,
                "gender": gender.lower(),
                "blood_type": None,
                "height_cm": None,
                "weight_kg": None
            },
            "medical_history": {
                "chronic_conditions": [],
                "allergies": [],
                "current_medications": [],
                "past_surgeries": [],
                "family_history": []
            },
            "lifestyle": {
                "smoking": None,
                "alcohol": None,
                "exercise_frequency": None,
                "diet_type": None
            },
            "tracking_preferences": {
                "track_womens_health": gender.lower() == "female",
                "track_mens_health": gender.lower() == "male"
            },
            "created_at": datetime.now(),
            "updated_at": datetime.now()
        }
        self.profiles.insert_one(profile)

    def get_profile(self, user_id: str) -> Optional[Dict]:
        """Get user profile"""
        try:
            return self.profiles.find_one({"user_id": user_id}, {"_id": 0})
        except Exception as e:
            print(f"Error fetching profile: {e}")
            return None

    def update_profile(self, user_id: str, updates: Dict) -> Dict:
        """Update user profile"""
        try:
            updates["updated_at"] = datetime.now()
            update_dict = {}
            
            for key, value in updates.items():
                if key in ["personal_info", "medical_history", "lifestyle"] and isinstance(value, dict):
                    for subk, subv in value.items():
                        update_dict[f"{key}.{subk}"] = subv
                else:
                    update_dict[key] = value
            
            res = self.profiles.update_one(
                {"user_id": user_id},
                {"$set": update_dict}
            )
            
            if res.modified_count > 0:
                return {"success": True, "message": "Profile updated successfully"}
            return {"success": True, "message": "No changes made"}
        except Exception as e:
            return {"success": False, "error": f"Profile update failed: {e}"}

    def save_conversation(self, user_id: str, query: str, response: Dict) -> str:
        """Save conversation history"""
        try:
            conv_id = str(uuid.uuid4())
            doc = {
                "conversation_id": conv_id,
                "user_id": user_id,
                "query": query,
                "response": response.get("answer", ""),
                "confidence": response.get("confidence", ""),
                "timestamp": datetime.now()
            }
            self.conversations.insert_one(doc)
            return conv_id
        except Exception as e:
            print(f"Error saving conversation: {e}")
            return ""

    def get_conversation_history(self, user_id: str, limit: int = 10) -> List[Dict]:
        """Get conversation history"""
        try:
            return list(self.conversations.find(
                {"user_id": user_id},
                {"_id": 0}
            ).sort("timestamp", DESCENDING).limit(limit))
        except Exception as e:
            print(f"Error fetching history: {e}")
            return []


class EnhancedRAGPipeline:
    """Enhanced RAG pipeline with Ollama integration"""
    
    def __init__(self, vectorstore, user_manager: UserManager, model_name: str = "llama3"):
        self.vectorstore = vectorstore
        self.user_manager = user_manager
        self.model_name = model_name
        
        if HAS_LANGCHAIN and vectorstore is not None:
            try:
                self.llm = Ollama(model=model_name, temperature=0.2)
                self.retriever = vectorstore.as_retriever(
                    search_type="mmr",
                    search_kwargs={"k": 6, "fetch_k": 12}
                )
                print(f"✓ RAG pipeline initialized with {model_name}")
            except Exception as e:
                print(f"Warning: Ollama initialization failed: {e}")
                self.llm = None
                self.retriever = None
        else:
            self.llm = None
            self.retriever = None
        
        self.prompt_template = """You are MedSage, an expert medical AI assistant. Provide helpful, accurate, and empathetic medical information.

PATIENT CONTEXT:
{patient_context}

RECENT MEDICAL REPORTS:
{report_context}

MEDICAL KNOWLEDGE BASE:
{medical_knowledge}

USER QUESTION: {question}

Please provide a comprehensive, personalized response considering the patient's context and medical history. If you're unsure about something, acknowledge it. Always recommend consulting with healthcare professionals for serious concerns.

RESPONSE:"""

    def query(self, user_id: str, question: str, report_ids: Optional[List[str]] = None) -> Dict:
        """Process user query with RAG"""
        try:
            # Retrieve medical knowledge
            medical_knowledge = ""
            if self.retriever is not None:
                try:
                    docs = self.retriever.invoke(question)
                    medical_knowledge = "\n\n".join([d.page_content for d in docs[:6]])
                except Exception as e:
                    print(f"Retrieval error: {e}")
                    medical_knowledge = "Unable to retrieve knowledge base"
            
            # Build context
            patient_context = self._build_patient_context(user_id, question)
            report_context = self._build_report_context(user_id, report_ids)
            
            # Format prompt
            prompt = self.prompt_template.format(
                patient_context=patient_context,
                report_context=report_context,
                medical_knowledge=medical_knowledge,
                question=question
            )
            
            # Generate response
            if self.llm is not None:
                try:
                    response = self.llm.invoke(prompt)
                    confidence = "high"
                except Exception as e:
                    print(f"LLM error: {e}")
                    response = self._generate_fallback_response(question, patient_context)
                    confidence = "low"
            else:
                response = self._generate_fallback_response(question, patient_context)
                confidence = "low"
            
            result = {
                "query": question,
                "answer": response,
                "confidence": confidence,
                "timestamp": datetime.now().isoformat(),
                "reports_used": len(report_ids or [])
            }
            
            # Save conversation
            try:
                self.user_manager.save_conversation(user_id, question, result)
            except Exception:
                pass
            
            return result
        except Exception as e:
            return {
                "query": question,
                "answer": f"I encountered an error processing your question: {e}. Please try rephrasing or contact support.",
                "confidence": "error",
                "timestamp": datetime.now().isoformat()
            }

    def _build_patient_context(self, user_id: str, query: str) -> str:
        """Build comprehensive patient context"""
        parts = []
        
        try:
            profile = self.user_manager.get_profile(user_id)
            if profile:
                personal = profile.get("personal_info", {})
                medical = profile.get("medical_history", {})
                
                parts.append(f"Age: {personal.get('age', 'Unknown')}")
                parts.append(f"Gender: {personal.get('gender', 'Unknown')}")
                
                if personal.get("blood_type"):
                    parts.append(f"Blood Type: {personal.get('blood_type')}")
                
                # Lifestyle summary
                lifestyle = self.user_manager.lifestyle_tracker.get_lifestyle_summary(user_id, 7)
                if lifestyle.get("success"):
                    parts.append("\nRecent Lifestyle (7 days):")
                    if lifestyle.get("avg_sleep_hours"):
                        parts.append(f"- Avg Sleep: {lifestyle.get('avg_sleep_hours')} hours")
                    if lifestyle.get("avg_exercise_minutes"):
                        parts.append(f"- Avg Exercise: {lifestyle.get('avg_exercise_minutes')} minutes")
                    if lifestyle.get("avg_stress"):
                        parts.append(f"- Avg Stress Level: {lifestyle.get('avg_stress')}/10")
                
                # Medical history
                if medical.get("chronic_conditions"):
                    parts.append(f"\nChronic Conditions: {', '.join(medical.get('chronic_conditions'))}")
                if medical.get("allergies"):
                    parts.append(f"Allergies: {', '.join(medical.get('allergies'))}")
        except Exception as e:
            parts.append(f"[Error loading context: {e}]")
        
        return "\n".join(parts) if parts else "No patient context available"

    def _build_report_context(self, user_id: str, report_ids: Optional[List[str]] = None) -> str:
        """Build report context"""
        parts = []
        
        try:
            if not report_ids:
                reports = self.user_manager.report_manager.get_user_reports(user_id)[:3]
            else:
                reports = []
                for rid in report_ids:
                    r = self.user_manager.report_manager.get_report_content(rid)
                    if r:
                        reports.append(r)
            
            if not reports:
                return "No medical reports available"
            
            for i, r in enumerate(reports, 1):
                parts.append(f"--- Report {i}: {r.get('report_name')} ({r.get('report_type')}) ---")
                text = (r.get("extracted_text") or "")[:1500]
                parts.append(text or "[No text extracted]")
        except Exception as e:
            parts.append(f"[Error loading reports: {e}]")
        
        return "\n".join(parts)

    def _generate_fallback_response(self, question: str, context: str) -> str:
        """Generate fallback response when LLM is unavailable"""
        return f"""I understand you're asking about: "{question}"

Based on your profile information:
{context}

I'm currently running in limited mode without access to the full AI model. Here's what I can suggest:

1. For specific medical concerns, please consult with a healthcare professional
2. You can upload medical reports for better tracking
3. Use the lifestyle tracking features to monitor your health trends
4. Schedule appointments with your healthcare providers

For detailed medical advice, please ensure the Ollama service is running with the llama3 model installed.

To install Ollama and the model:
1. Install Ollama from https://ollama.ai
2. Run: ollama pull llama3
3. Restart MedSage

Is there anything else I can help you with using the available features?"""


class MedicalRAGSystem:
    """Main medical RAG system"""
    
    def __init__(self, data_dir: str = "./medical_data", persist_dir: str = "./medical_knowledge",
                 mongodb_uri: str = "mongodb://localhost:27017/", model_name: str = "llama3"):
        self.data_dir = data_dir
        self.persist_dir = persist_dir
        self.model_name = model_name
        self.current_user = None
        
        print("=" * 60)
        print("MedSage - Medical RAG System")
        print("=" * 60)
        
        # Initialize MongoDB
        if not HAS_MONGODB:
            raise ImportError("MongoDB required. Install: pip install pymongo")
        
        try:
            self.user_manager = UserManager(mongodb_uri, database_name="medical_rag_system")
        except Exception as e:
            print(f"Error: Could not connect to MongoDB: {e}")
            print("\nPlease ensure MongoDB is running:")
            print("  - Install MongoDB from https://www.mongodb.com/")
            print("  - Start service: sudo systemctl start mongod")
            raise
        
        self.rag_pipeline = None

    def initialize_system(self, force_rebuild: bool = False):
        """Initialize the RAG system"""
        print("\nInitializing system...")
        
        vs = None
        if HAS_LANGCHAIN:
            mv = MedicalVectorStore(self.persist_dir)
            
            if os.path.exists(self.persist_dir) and not force_rebuild:
                print("Loading existing knowledge base...")
                vs = mv.load_vectorstore()
                if not vs:
                    print("Failed to load, rebuilding...")
                    force_rebuild = True
            
            if force_rebuild or not os.path.exists(self.persist_dir):
                print("Building knowledge base from CSV files...")
                processor = CSVDataProcessor(self.data_dir)
                docs = processor.process_all_files()
                
                if docs:
                    vs = mv.create_vectorstore(docs)
                else:
                    print("Warning: No documents found in data directory")
        else:
            print("Warning: LangChain not available, running without vector store")
        
        self.rag_pipeline = EnhancedRAGPipeline(vs, self.user_manager, model_name=self.model_name)
        print("✓ System initialization complete\n")

    def run(self):
        """Main application loop"""
        while True:
            try:
                if self.current_user is None:
                    self._show_welcome()
                else:
                    self._main_menu()
            except KeyboardInterrupt:
                print("\n\nGoodbye!")
                break
            except Exception as e:
                print(f"\nError: {e}")
                input("Press Enter to continue...")

    def _show_welcome(self):
        """Show welcome screen"""
        print("\n" + "=" * 60)
        print("MEDSAGE - Your Personal Medical Assistant")
        print("=" * 60)
        print("\n1. Sign Up (New User)")
        print("2. Login")
        print("3. Exit")
        print("-" * 60)
        
        choice = safe_input("Enter choice (1-3): ", "3")
        
        if choice == "1":
            self._signup()
        elif choice == "2":
            self._login()
        elif choice == "3":
            print("\nThank you for using MedSage!")
            exit(0)

    def _signup(self):
        """User registration"""
        print("\n" + "=" * 60)
        print("USER REGISTRATION")
        print("=" * 60)
        
        email = safe_input("Email: ")
        if not email:
            print("Email cannot be empty")
            return
        
        password = getpass.getpass("Password (min 6 characters): ")
        if not password:
            print("Password cannot be empty")
            return
        
        confirm = getpass.getpass("Confirm Password: ")
        if password != confirm:
            print("❌ Passwords don't match!")
            return
        
        name = safe_input("Full Name: ")
        age = safe_int_input("Age: ", 0, 1, 120)
        
        print("\nGender Options:")
        print("1. Male")
        print("2. Female")
        print("3. Other")
        gender_choice = safe_input("Select (1-3): ", "3")
        gender = {"1": "male", "2": "female", "3": "other"}.get(gender_choice, "other")
        
        result = self.user_manager.signup(email, password, name, age, gender)
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _login(self):
        """User login"""
        print("\n" + "=" * 60)
        print("USER LOGIN")
        print("=" * 60)
        
        email = safe_input("Email: ")
        if not email:
            return
        
        password = getpass.getpass("Password: ")
        if not password:
            return
        
        result = self.user_manager.login(email, password)
        
        if result.get("success"):
            self.current_user = result
            print(f"\n✓ Welcome, {result.get('full_name')}!")
            input("\nPress Enter to continue...")
        else:
            print(f"\n❌ {result.get('error')}")
            input("\nPress Enter to continue...")

    def _main_menu(self):
        """Main menu after login"""
        profile = self.user_manager.get_profile(self.current_user["user_id"])
        prefs = profile.get("tracking_preferences", {}) if profile else {}
        gender = profile.get("personal_info", {}).get("gender", "") if profile else ""
        
        print("\n" + "=" * 60)
        print(f"Welcome, {self.current_user.get('full_name')}!")
        print("=" * 60)
        print("\n📋 MAIN MENU")
        print("-" * 60)
        print("1.  💬 Ask MedSage (AI Consultation)")
        print("2.  📄 Upload Medical Report")
        print("3.  📁 View My Reports")
        print("4.  📊 Log Daily Lifestyle")
        print("5.  📈 View Lifestyle Summary")
        print("6.  💧 Log Water Intake")
        print("7.  🍽️  Log Meal")
        print("8.  💊 Medication Management")
        print("9.  📅 Appointment Scheduler")
        print("10. 👤 View/Update Profile")
        
        if gender == "female":
            print("11. 🌸 Women's Health Tracking")
        elif gender == "male":
            print("11. 💪 Men's Health Tracking")
        
        print("12. 📜 View Conversation History")
        print("13. 🚪 Logout")
        print("-" * 60)
        
        choice = safe_input("Enter choice: ")
        
        menu_actions = {
            "1": self._consult,
            "2": self._upload_report,
            "3": self._view_reports,
            "4": self._log_lifestyle,
            "5": self._view_lifestyle,
            "6": self._log_water,
            "7": self._log_meal,
            "8": self._medication_menu,
            "9": self._appointment_menu,
            "10": self._profile_menu,
            "11": self._gender_specific_menu,
            "12": self._view_history,
            "13": self._logout
        }
        
        action = menu_actions.get(choice)
        if action:
            action()
        else:
            print("\n❌ Invalid choice")
            input("\nPress Enter to continue...")

    def _consult(self):
        """AI consultation"""
        print("\n" + "=" * 60)
        print("AI MEDICAL CONSULTATION")
        print("=" * 60)
        
        question = safe_input("\nYour question: ")
        if not question:
            return
        
        print("\n⏳ Processing your question...")
        
        uid = self.current_user["user_id"]
        result = self.rag_pipeline.query(uid, question) if self.rag_pipeline else {
            "answer": "RAG pipeline unavailable"
        }
        
        print("\n" + "=" * 60)
        print("MEDSAGE RESPONSE")
        print("=" * 60)
        print(f"\n{result.get('answer')}")
        print(f"\nConfidence: {result.get('confidence')}")
        print("=" * 60)
        
        input("\nPress Enter to continue...")

    def _upload_report(self):
        """Upload medical report"""
        print("\n" + "=" * 60)
        print("UPLOAD MEDICAL REPORT")
        print("=" * 60)
        
        path = safe_input("\nFile path: ")
        if not path or not Path(path).exists():
            print("❌ File not found")
            input("\nPress Enter to continue...")
            return
        
        name = safe_input("Report name (or press Enter for filename): ") or Path(path).stem
        
        print("\nReport Types:")
        print("1. Blood Test")
        print("2. Imaging (X-ray, MRI, CT)")
        print("3. Prescription")
        print("4. Consultation Notes")
        print("5. Other")
        
        type_choice = safe_input("Select type (1-5): ", "5")
        rtype = {
            "1": "blood_test",
            "2": "imaging",
            "3": "prescription",
            "4": "consultation",
            "5": "other"
        }.get(type_choice, "other")
        
        notes = safe_input("Additional notes (optional): ")
        
        print("\n⏳ Uploading and processing file...")
        
        uid = self.current_user["user_id"]
        result = self.user_manager.report_manager.upload_report(uid, path, rtype, name, notes)
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
            print(f"Text extracted: {result.get('extracted_text_length', 0)} characters")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _view_reports(self):
        """View medical reports"""
        print("\n" + "=" * 60)
        print("YOUR MEDICAL REPORTS")
        print("=" * 60)
        
        uid = self.current_user["user_id"]
        reports = self.user_manager.report_manager.get_user_reports(uid)
        
        if not reports:
            print("\nNo reports found")
        else:
            print(f"\nTotal Reports: {len(reports)}\n")
            for i, r in enumerate(reports, 1):
                print(f"{i}. {r.get('report_name')}")
                print(f"   Type: {r.get('report_type')}")
                print(f"   Date: {r.get('upload_date')}")
                print(f"   ID: {r.get('report_id')}")
                print()
        
        input("\nPress Enter to continue...")

    def _log_lifestyle(self):
        """Log daily lifestyle data"""
        print("\n" + "=" * 60)
        print("LOG DAILY LIFESTYLE")
        print("=" * 60)
        
        date = datetime.now().date().isoformat()
        print(f"\nDate: {date}")
        
        sleep = safe_float_input("\nSleep hours (0-24): ", 7.0, 0, 24)
        exercise = safe_int_input("Exercise minutes: ", 0, 0, 1440)
        
        print("\nMood Scale (1-10):")
        print("1-3: Poor | 4-6: Average | 7-9: Good | 10: Excellent")
        mood = safe_int_input("Mood rating: ", 5, 1, 10)
        
        print("\nStress Level (1-10):")
        print("1-3: Low | 4-6: Moderate | 7-9: High | 10: Very High")
        stress = safe_int_input("Stress level: ", 5, 1, 10)
        
        data = {
            "sleep_hours": sleep,
            "exercise_minutes": exercise,
            "mood": mood,
            "stress_level": stress
        }
        
        uid = self.current_user["user_id"]
        result = self.user_manager.lifestyle_tracker.log_daily_lifestyle(uid, date, data)
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
            print("\nSummary:")
            print(f"  Sleep: {sleep} hours")
            print(f"  Exercise: {exercise} minutes")
            print(f"  Mood: {mood}/10")
            print(f"  Stress: {stress}/10")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _view_lifestyle(self):
        """View lifestyle summary"""
        print("\n" + "=" * 60)
        print("LIFESTYLE SUMMARY")
        print("=" * 60)
        
        print("\nPeriod Options:")
        print("1. Last 7 days")
        print("2. Last 30 days")
        print("3. Last 90 days")
        
        choice = safe_input("Select (1-3): ", "1")
        days = {"1": 7, "2": 30, "3": 90}.get(choice, 7)
        
        uid = self.current_user["user_id"]
        result = self.user_manager.lifestyle_tracker.get_lifestyle_summary(uid, days)
        
        if result.get("success"):
            print(f"\n📊 Summary for Last {days} Days")
            print("-" * 60)
            print(f"Total Entries: {result.get('total_entries', 0)}")
            
            if result.get('total_entries', 0) > 0:
                print(f"\nAverages:")
                if result.get('avg_sleep_hours'):
                    print(f"  Sleep: {result.get('avg_sleep_hours')} hours")
                if result.get('avg_exercise_minutes'):
                    print(f"  Exercise: {result.get('avg_exercise_minutes')} minutes")
                if result.get('avg_mood'):
                    print(f"  Mood: {result.get('avg_mood')}/10")
                if result.get('avg_stress'):
                    print(f"  Stress: {result.get('avg_stress')}/10")
            else:
                print("\n" + result.get('message', 'No data available'))
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _log_water(self):
        """Log water intake"""
        print("\n" + "=" * 60)
        print("LOG WATER INTAKE")
        print("=" * 60)
        
        date = datetime.now().date().isoformat()
        print(f"\nDate: {date}")
        
        print("\nQuick Options:")
        print("1. Glass (250ml)")
        print("2. Bottle (500ml)")
        print("3. Large Bottle (1000ml)")
        print("4. Custom amount")
        
        choice = safe_input("Select (1-4): ", "1")
        
        if choice == "4":
            amount = safe_int_input("Amount (ml): ", 250, 1, 5000)
        else:
            amount = {"1": 250, "2": 500, "3": 1000}.get(choice, 250)
        
        uid = self.current_user["user_id"]
        result = self.user_manager.lifestyle_tracker.log_water_intake(uid, date, amount)
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
            print(f"Logged: {amount}ml")
            
            # Show daily total
            daily = self.user_manager.lifestyle_tracker.get_daily_water_intake(uid, date)
            if daily.get("success"):
                print(f"\nToday's Total: {daily.get('total_ml')}ml / {daily.get('target_ml')}ml")
                print(f"Progress: {daily.get('percentage', 0):.1f}%")
                print(f"Glasses (8oz): {daily.get('glasses_8oz')}")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _log_meal(self):
        """Log meal"""
        print("\n" + "=" * 60)
        print("LOG MEAL")
        print("=" * 60)
        
        date = datetime.now().date().isoformat()
        
        print("\nMeal Type:")
        print("1. Breakfast")
        print("2. Lunch")
        print("3. Dinner")
        print("4. Snack")
        
        choice = safe_input("Select (1-4): ", "1")
        meal_type = {"1": "breakfast", "2": "lunch", "3": "dinner", "4": "snack"}.get(choice, "breakfast")
        
        description = safe_input("\nDescribe your meal: ")
        if not description:
            print("❌ Description required")
            input("\nPress Enter to continue...")
            return
        
        calories = safe_int_input("Estimated calories (optional, 0 to skip): ", 0, 0, 5000)
        calories = calories if calories > 0 else None
        
        uid = self.current_user["user_id"]
        result = self.user_manager.lifestyle_tracker.log_meal(uid, date, meal_type, description, calories)
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _medication_menu(self):
        """Medication management menu"""
        while True:
            print("\n" + "=" * 60)
            print("MEDICATION MANAGEMENT")
            print("=" * 60)
            print("\n1. Add Medication")
            print("2. View Active Medications")
            print("3. Log Dose Taken")
            print("4. Back to Main Menu")
            print("-" * 60)
            
            choice = safe_input("Enter choice (1-4): ")
            
            if choice == "1":
                self._add_medication()
            elif choice == "2":
                self._view_medications()
            elif choice == "3":
                self._log_medication_dose()
            elif choice == "4":
                break

    def _add_medication(self):
        """Add new medication"""
        print("\n--- Add Medication ---")
        
        name = safe_input("Medication name: ")
        if not name:
            print("❌ Name required")
            return
        
        dosage = safe_input("Dosage (e.g., 500mg, 2 tablets): ")
        frequency = safe_input("Frequency (e.g., twice daily, every 8 hours): ")
        start_date = safe_input(f"Start date (YYYY-MM-DD) or press Enter for today: ") or datetime.now().date().isoformat()
        end_date = safe_input("End date (optional, YYYY-MM-DD): ")
        notes = safe_input("Notes (optional): ")
        
        uid = self.current_user["user_id"]
        result = self.user_manager.medication_manager.add_medication(
            uid, name, dosage, frequency, start_date, end_date, notes
        )
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _view_medications(self):
        """View active medications"""
        print("\n--- Active Medications ---")
        
        uid = self.current_user["user_id"]
        result = self.user_manager.medication_manager.get_active_medications(uid)
        
        if result.get("success"):
            meds = result.get("medications", [])
            if not meds:
                print("\nNo active medications")
            else:
                print(f"\nTotal: {result.get('count')}\n")
                for i, med in enumerate(meds, 1):
                    print(f"{i}. {med.get('name')}")
                    print(f"   Dosage: {med.get('dosage')}")
                    print(f"   Frequency: {med.get('frequency')}")
                    print(f"   Started: {med.get('start_date')}")
                    if med.get('end_date'):
                        print(f"   Ends: {med.get('end_date')}")
                    print(f"   ID: {med.get('medication_id')}")
                    print()
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _log_medication_dose(self):
        """Log medication dose"""
        print("\n--- Log Medication Dose ---")
        
        # First show medications
        uid = self.current_user["user_id"]
        result = self.user_manager.medication_manager.get_active_medications(uid)
        
        if not result.get("success") or not result.get("medications"):
            print("\nNo active medications found. Add a medication first.")
            input("\nPress Enter to continue...")
            return
        
        meds = result.get("medications", [])
        print("\nYour Medications:")
        for i, med in enumerate(meds, 1):
            print(f"{i}. {med.get('name')} - {med.get('dosage')}")
        
        choice = safe_int_input(f"\nSelect medication (1-{len(meds)}): ", 1, 1, len(meds))
        selected_med = meds[choice - 1]
        
        date = datetime.now().date().isoformat()
        time = datetime.now().strftime("%H:%M")
        
        taken = safe_input(f"Did you take it? (y/n, default yes): ", "y").lower() == "y"
        
        result = self.user_manager.medication_manager.log_dose(
            uid, selected_med["medication_id"], date, time, taken
        )
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _appointment_menu(self):
        """Appointment management menu"""
        while True:
            print("\n" + "=" * 60)
            print("APPOINTMENT SCHEDULER")
            print("=" * 60)
            print("\n1. Schedule Appointment")
            print("2. View Upcoming Appointments")
            print("3. Back to Main Menu")
            print("-" * 60)
            
            choice = safe_input("Enter choice (1-3): ")
            
            if choice == "1":
                self._schedule_appointment()
            elif choice == "2":
                self._view_appointments()
            elif choice == "3":
                break

    def _schedule_appointment(self):
        """Schedule new appointment"""
        print("\n--- Schedule Appointment ---")
        
        doctor = safe_input("Doctor/Specialist name: ")
        if not doctor:
            print("❌ Doctor name required")
            return
        
        date = safe_input("Date (YYYY-MM-DD): ")
        time = safe_input("Time (HH:MM): ")
        purpose = safe_input("Purpose of visit: ")
        location = safe_input("Location (optional): ")
        notes = safe_input("Notes (optional): ")
        
        uid = self.current_user["user_id"]
        result = self.user_manager.appointment_scheduler.schedule_appointment(
            uid, date, time, doctor, purpose, location, notes
        )
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _view_appointments(self):
        """View upcoming appointments"""
        print("\n--- Upcoming Appointments ---")
        
        uid = self.current_user["user_id"]
        result = self.user_manager.appointment_scheduler.get_upcoming_appointments(uid, 90)
        
        if result.get("success"):
            appts = result.get("appointments", [])
            if not appts:
                print("\nNo upcoming appointments")
            else:
                print(f"\nTotal: {result.get('count')}\n")
                for i, appt in enumerate(appts, 1):
                    print(f"{i}. {appt.get('doctor')}")
                    print(f"   Date: {appt.get('date')} at {appt.get('time')}")
                    print(f"   Purpose: {appt.get('purpose')}")
                    if appt.get('location'):
                        print(f"   Location: {appt.get('location')}")
                    print()
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _profile_menu(self):
        """Profile management menu"""
        while True:
            print("\n" + "=" * 60)
            print("PROFILE MANAGEMENT")
            print("=" * 60)
            print("\n1. View Profile")
            print("2. Update Personal Info")
            print("3. Update Medical History")
            print("4. Back to Main Menu")
            print("-" * 60)
            
            choice = safe_input("Enter choice (1-4): ")
            
            if choice == "1":
                self._view_profile()
            elif choice == "2":
                self._update_personal_info()
            elif choice == "3":
                self._update_medical_history()
            elif choice == "4":
                break

    def _view_profile(self):
        """View user profile"""
        print("\n--- Your Profile ---")
        
        uid = self.current_user["user_id"]
        profile = self.user_manager.get_profile(uid)
        
        if profile:
            personal = profile.get("personal_info", {})
            medical = profile.get("medical_history", {})
            lifestyle = profile.get("lifestyle", {})
            
            print("\nPersonal Information:")
            print(f"  Age: {personal.get('age')}")
            print(f"  Gender: {personal.get('gender')}")
            print(f"  Blood Type: {personal.get('blood_type') or 'Not set'}")
            print(f"  Height: {personal.get('height_cm') or 'Not set'} cm")
            print(f"  Weight: {personal.get('weight_kg') or 'Not set'} kg")
            
            print("\nMedical History:")
            print(f"  Chronic Conditions: {', '.join(medical.get('chronic_conditions', [])) or 'None'}")
            print(f"  Allergies: {', '.join(medical.get('allergies', [])) or 'None'}")
            print(f"  Current Medications: {', '.join(medical.get('current_medications', [])) or 'None'}")
            
            print("\nLifestyle:")
            print(f"  Smoking: {lifestyle.get('smoking') or 'Not set'}")
            print(f"  Alcohol: {lifestyle.get('alcohol') or 'Not set'}")
            print(f"  Exercise: {lifestyle.get('exercise_frequency') or 'Not set'}")
            print(f"  Diet: {lifestyle.get('diet_type') or 'Not set'}")
        else:
            print("\n❌ Profile not found")
        
        input("\nPress Enter to continue...")

    def _update_personal_info(self):
        """Update personal information"""
        print("\n--- Update Personal Info ---")
        
        blood_type = safe_input("Blood Type (A+, B+, O-, etc.) or press Enter to skip: ")
        height = safe_float_input("Height (cm) or 0 to skip: ", 0, 0, 300)
        weight = safe_float_input("Weight (kg) or 0 to skip: ", 0, 0, 500)
        
        updates = {"personal_info": {}}
        if blood_type:
            updates["personal_info"]["blood_type"] = blood_type
        if height > 0:
            updates["personal_info"]["height_cm"] = height
        if weight > 0:
            updates["personal_info"]["weight_kg"] = weight
        
        if updates["personal_info"]:
            uid = self.current_user["user_id"]
            result = self.user_manager.update_profile(uid, updates)
            
            if result.get("success"):
                print(f"\n✓ {result.get('message')}")
            else:
                print(f"\n❌ {result.get('error')}")
        else:
            print("\n❌ No updates provided")
        
        input("\nPress Enter to continue...")

    def _update_medical_history(self):
        """Update medical history"""
        print("\n--- Update Medical History ---")
        
        print("\nAdd chronic conditions (comma-separated, or press Enter to skip):")
        conditions = safe_input("Conditions: ")
        
        print("\nAdd allergies (comma-separated, or press Enter to skip):")
        allergies = safe_input("Allergies: ")
        
        updates = {"medical_history": {}}
        
        if conditions:
            cond_list = [c.strip() for c in conditions.split(",") if c.strip()]
            updates["medical_history"]["chronic_conditions"] = cond_list
        
        if allergies:
            allergy_list = [a.strip() for a in allergies.split(",") if a.strip()]
            updates["medical_history"]["allergies"] = allergy_list
        
        if updates["medical_history"]:
            uid = self.current_user["user_id"]
            result = self.user_manager.update_profile(uid, updates)
            
            if result.get("success"):
                print(f"\n✓ {result.get('message')}")
            else:
                print(f"\n❌ {result.get('error')}")
        else:
            print("\n❌ No updates provided")
        
        input("\nPress Enter to continue...")

    def _gender_specific_menu(self):
        """Gender-specific health tracking"""
        uid = self.current_user["user_id"]
        profile = self.user_manager.get_profile(uid)
        gender = profile.get("personal_info", {}).get("gender", "") if profile else ""
        
        if gender == "female":
            self._womens_health_menu()
        elif gender == "male":
            self._mens_health_menu()
        else:
            print("\n❌ Gender-specific tracking not available for your profile")
            input("\nPress Enter to continue...")

    def _womens_health_menu(self):
        """Women's health tracking menu"""
        while True:
            print("\n" + "=" * 60)
            print("WOMEN'S HEALTH TRACKING")
            print("=" * 60)
            print("\n1. Log Period")
            print("2. Predict Next Period")
            print("3. Log Symptom")
            print("4. Back to Main Menu")
            print("-" * 60)
            
            choice = safe_input("Enter choice (1-4): ")
            
            if choice == "1":
                self._log_period()
            elif choice == "2":
                self._predict_period()
            elif choice == "3":
                self._log_symptom()
            elif choice == "4":
                break

    def _log_period(self):
        """Log menstrual period"""
        print("\n--- Log Period ---")
        
        start_date = safe_input(f"Start date (YYYY-MM-DD) or press Enter for today: ") or datetime.now().date().isoformat()
        end_date = safe_input("End date (YYYY-MM-DD, optional): ")
        
        print("\nFlow intensity:")
        print("1. Light")
        print("2. Medium")
        print("3. Heavy")
        
        flow_choice = safe_input("Select (1-3): ", "2")
        flow = {"1": "light", "2": "medium", "3": "heavy"}.get(flow_choice, "medium")
        
        notes = safe_input("Notes (optional): ")
        
        uid = self.current_user["user_id"]
        result = self.user_manager.womens_health.log_period(uid, start_date, end_date, flow, notes)
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _predict_period(self):
        """Predict next period"""
        print("\n--- Period Prediction ---")
        
        uid = self.current_user["user_id"]
        result = self.user_manager.womens_health.predict_next_period(uid)
        
        if result.get("success"):
            print("\n📅 Prediction Results:")
            print(f"  Next Period: {result.get('next_period_date')}")
            print(f"  Confidence: {result.get('confidence')}")
            print(f"  Average Cycle: {result.get('average_cycle_length')} days")
            print(f"  Cycle Variation: {result.get('cycle_variation')} days")
            print(f"\n  Ovulation Date: {result.get('ovulation_date')}")
            print(f"  Fertile Window: {result.get('fertile_window_start')} to {result.get('fertile_window_end')}")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _log_symptom(self):
        """Log health symptom"""
        print("\n--- Log Symptom ---")
        
        print("\nCommon symptoms:")
        print("1. Cramps")
        print("2. Headache")
        print("3. Bloating")
        print("4. Fatigue")
        print("5. Mood changes")
        print("6. Other")
        
        choice = safe_input("Select (1-6): ", "1")
        symptom_types = {
            "1": "cramps",
            "2": "headache",
            "3": "bloating",
            "4": "fatigue",
            "5": "mood_changes",
            "6": safe_input("Describe symptom: ")
        }
        symptom = symptom_types.get(choice, "other")
        
        print("\nSeverity (1-10):")
        severity = safe_int_input("Rating: ", 5, 1, 10)
        
        notes = safe_input("Notes (optional): ")
        date = datetime.now().date().isoformat()
        
        uid = self.current_user["user_id"]
        result = self.user_manager.womens_health.log_symptom(uid, date, symptom, severity, notes)
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _mens_health_menu(self):
        """Men's health tracking menu"""
        while True:
            print("\n" + "=" * 60)
            print("MEN'S HEALTH TRACKING")
            print("=" * 60)
            print("\n1. Log Daily Health Metrics")
            print("2. View Health Insights")
            print("3. Back to Main Menu")
            print("-" * 60)
            
            choice = safe_input("Enter choice (1-3): ")
            
            if choice == "1":
                self._log_mens_health()
            elif choice == "2":
                self._view_mens_insights()
            elif choice == "3":
                break

    def _log_mens_health(self):
        """Log men's health metrics"""
        print("\n--- Log Health Metrics ---")
        
        date = datetime.now().date().isoformat()
        
        print("\nEnergy Level (1-10):")
        energy = safe_int_input("Rating: ", 5, 1, 10)
        
        print("\nSleep Quality (1-10):")
        sleep_quality = safe_int_input("Rating: ", 5, 1, 10)
        
        print("\nPhysical Activity (1-10):")
        activity = safe_int_input("Rating: ", 5, 1, 10)
        
        data = {
            "energy_level": energy,
            "sleep_quality": sleep_quality,
            "physical_activity": activity
        }
        
        uid = self.current_user["user_id"]
        result = self.user_manager.mens_health.log_daily_health(uid, date, data)
        
        if result.get("success"):
            print(f"\n✓ {result.get('message')}")
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _view_mens_insights(self):
        """View men's health insights"""
        print("\n--- Health Insights ---")
        
        print("\nPeriod Options:")
        print("1. Last 7 days")
        print("2. Last 30 days")
        print("3. Last 90 days")
        
        choice = safe_input("Select (1-3): ", "2")
        days = {"1": 7, "2": 30, "3": 90}.get(choice, 30)
        
        uid = self.current_user["user_id"]
        result = self.user_manager.mens_health.get_health_insights(uid, days)
        
        if result.get("success"):
            print(f"\n📊 Insights for Last {days} Days")
            print("-" * 60)
            print(f"Total Entries: {result.get('total_entries', 0)}")
            
            if result.get('total_entries', 0) > 0:
                print(f"\nAverages:")
                if result.get('avg_energy'):
                    print(f"  Energy Level: {result.get('avg_energy')}/10")
                if result.get('avg_sleep_quality'):
                    print(f"  Sleep Quality: {result.get('avg_sleep_quality')}/10")
                print(f"\nTrend: {result.get('trend')}")
            else:
                print("\n" + result.get('message', 'No data available'))
        else:
            print(f"\n❌ {result.get('error')}")
        
        input("\nPress Enter to continue...")

    def _view_history(self):
        """View conversation history"""
        print("\n" + "=" * 60)
        print("CONVERSATION HISTORY")
        print("=" * 60)
        
        uid = self.current_user["user_id"]
        history = self.user_manager.get_conversation_history(uid, 10)
        
        if not history:
            print("\nNo conversation history")
        else:
            print(f"\nLast {len(history)} conversations:\n")
            for i, conv in enumerate(history, 1):
                print(f"{i}. Q: {conv.get('query')}")
                print(f"   A: {conv.get('response')[:100]}...")
                print(f"   Time: {conv.get('timestamp')}")
                print()
        
        input("\nPress Enter to continue...")

    def _logout(self):
        """Logout user"""
        print("\n👋 Logging out...")
        self.current_user = None
        input("\nPress Enter to continue...")


# CLI Entry Point
if __name__ == "__main__":
    import argparse
    
    parser = argparse.ArgumentParser(
        description="MedSage - Advanced Medical RAG System",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python MedSage_2.py
  python MedSage_2.py --data-dir ./my_medical_data
  python MedSage_2.py --model llama2 --force-rebuild
  
Requirements:
  - MongoDB running on localhost:27017
  - Ollama installed with llama3 model (optional but recommended)
  - Python packages: pymongo, pandas, numpy, langchain-community, chromadb
        """
    )
    
    parser.add_argument("--data-dir", default="./medical_data",
                       help="Directory containing medical CSV data files")
    parser.add_argument("--persist-dir", default="./medical_knowledge",
                       help="Directory for vector store persistence")
    parser.add_argument("--mongodb-uri", default="mongodb://localhost:27017/",
                       help="MongoDB connection URI")
    parser.add_argument("--model", default="llama3",
                       help="Ollama model name (e.g., llama3, llama2, mistral)")
    parser.add_argument("--force-rebuild", action="store_true",
                       help="Force rebuild of vector store")
    
    args = parser.parse_args()
    
    print("\n" + "=" * 60)
    print("Starting MedSage Medical RAG System")
    print("=" * 60)
    
    # Check dependencies
    print("\nChecking dependencies...")
    missing_deps = []
    
    if not HAS_MONGODB:
        missing_deps.append("pymongo")
    if not HAS_LANGCHAIN:
        missing_deps.append("langchain-community")
    
    if missing_deps:
        print(f"\n⚠️  Warning: Missing optional dependencies: {', '.join(missing_deps)}")
        print("Install with: pip install " + " ".join(missing_deps))
        print("\nSome features may be limited.\n")
    
    print("\nDependency Status:")
    print(f"  ✓ MongoDB: {'Available' if HAS_MONGODB else '❌ Not installed'}")
    print(f"  ✓ LangChain: {'Available' if HAS_LANGCHAIN else '❌ Not installed'}")
    print(f"  ✓ PDF Processing: {'Available' if HAS_PDF else '❌ Not installed'}")
    print(f"  ✓ OCR: {'Available' if HAS_OCR else '❌ Not installed'}")
    print(f"  ✓ DOCX: {'Available' if HAS_DOCX else '❌ Not installed'}")
    print(f"  ✓ Plotting: {'Available' if HAS_PLOTTING else '❌ Not installed'}")
    
    try:
        # Initialize system
        system = MedicalRAGSystem(
            data_dir=args.data_dir,
            persist_dir=args.persist_dir,
            mongodb_uri=args.mongodb_uri,
            model_name=args.model
        )
        
        # Initialize knowledge base
        system.initialize_system(force_rebuild=args.force_rebuild)
        
        # Run main application
        system.run()
        
    except KeyboardInterrupt:
        print("\n\n👋 Goodbye! Stay healthy!")
    except ConnectionError as e:
        print(f"\n❌ Connection Error: {e}")
        print("\nPlease ensure MongoDB is running:")
        print("  Ubuntu/Debian: sudo systemctl start mongod")
        print("  macOS: brew services start mongodb-community")
        print("  Windows: net start MongoDB")
    except ImportError as e:
        print(f"\n❌ Import Error: {e}")
        print("\nPlease install required dependencies:")
        print("  pip install pymongo pandas numpy")
        print("\nOptional but recommended:")
        print("  pip install langchain-community chromadb sentence-transformers")
        print("  pip install PyPDF2 python-docx pillow pytesseract")
    except Exception as e:
        print(f"\n❌ Unexpected Error: {e}")
        print("\nPlease check your configuration and try again.")
        import traceback
        traceback.print_exc()
    finally:
        print("\nThank you for using MedSage!")
        print("=" * 60)